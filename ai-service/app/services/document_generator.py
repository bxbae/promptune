from __future__ import annotations

import re
import textwrap
from io import BytesIO

import fitz
from docx import Document as DocxDocument


def _safe_filename(title: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', "_", title).strip()
    return name or "document"


def _generate_docx(title: str, content: str) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=0)

    for line in content.splitlines():
        if line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _generate_pdf(title: str, content: str) -> bytes:
    pdf = fitz.open()

    page_width = 595
    page_height = 842
    margin = 50
    bottom_margin = 50

    page = pdf.new_page(width=page_width, height=page_height)
    y = 60

    def new_page():
        nonlocal page, y
        page = pdf.new_page(width=page_width, height=page_height)
        y = 60

    for title_line in textwrap.wrap(title, width=32) or [title]:
        page.insert_text(
            (margin, y),
            title_line,
            fontsize=17,
            fontname="korea",
        )
        y += 24

    y += 12

    for paragraph in content.splitlines():
        if not paragraph.strip():
            y += 10
            continue

        lines = textwrap.wrap(
            paragraph,
            width=48,
            break_long_words=True,
            replace_whitespace=False,
        ) or [""]

        for line in lines:
            if y >= page_height - bottom_margin:
                new_page()

            page.insert_text(
                (margin, y),
                line,
                fontsize=10.5,
                fontname="korea",
            )
            y += 17

        y += 4

    result = pdf.tobytes()
    pdf.close()
    return result



def _generate_txt(title: str, content: str) -> bytes:
    text = f"{title}\n\n{content}\n"
    return text.encode("utf-8")


def _generate_md(title: str, content: str) -> bytes:
    markdown = f"# {title}\n\n{content}\n"
    return markdown.encode("utf-8")


def generate_document(
    title: str,
    content: str,
    output_format: str,
) -> tuple[bytes, str, str]:
    fmt = output_format.lower().strip()
    safe_title = _safe_filename(title)

    if fmt == "docx":
        return (
            _generate_docx(title, content),
            f"{safe_title}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if fmt == "pdf":
        return (
            _generate_pdf(title, content),
            f"{safe_title}.pdf",
            "application/pdf",
        )

    if fmt == "txt":
        return (
            _generate_txt(title, content),
            f"{safe_title}.txt",
            "text/plain; charset=utf-8",
        )

    if fmt == "md":
        return (
            _generate_md(title, content),
            f"{safe_title}.md",
            "text/markdown; charset=utf-8",
        )

    raise ValueError("지원 형식은 pdf, docx, txt 또는 md입니다.")
