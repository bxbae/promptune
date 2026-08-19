from __future__ import annotations

import io
from pathlib import Path

import numpy as np
from docx import Document as DocxDocument
import pymupdf

from app.services.retrieval.chunker import chunk_text
from app.services.retrieval.rag_retriever import (
    EXPECTED_DIM,
    get_connection,
    get_model,
    vector_literal,
)


SUPPORTED_FILE_TYPES = {"pdf", "docx", "txt", "md"}
MAX_FILE_SIZE = 20 * 1024 * 1024


def resolve_file_type(filename: str | None, file_type: str | None) -> str:
    if file_type:
        value = file_type.lower().strip().lstrip(".")
    elif filename:
        value = Path(filename).suffix.lower().lstrip(".")
    else:
        raise ValueError("file_type 또는 filename이 필요합니다.")

    if value not in SUPPORTED_FILE_TYPES:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {value}")

    return value


def extract_pdf_text(file_bytes: bytes) -> str:
    """PDF에서 텍스트를 추출한다."""
    try:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        pages = []

        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                pages.append(page_text.strip())

        return "\n".join(pages).strip()

    except Exception as exc:
        raise ValueError(f"PDF 텍스트 추출에 실패했습니다: {exc}") from exc


def extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def decode_text_file(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            pass

    raise ValueError("TXT/MD 인코딩을 읽을 수 없습니다.")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    if not file_bytes:
        raise ValueError("파일이 비어 있습니다.")

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError("파일 크기가 20MB를 초과했습니다.")

    if file_type == "pdf":
        text = extract_pdf_text(file_bytes)
    elif file_type == "docx":
        text = extract_docx_text(file_bytes)
    elif file_type in {"txt", "md"}:
        text = decode_text_file(file_bytes)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {file_type}")

    if not text.strip():
        raise ValueError("문서에서 텍스트를 추출하지 못했습니다.")

    return text.strip()


def embed_chunks(chunks: list[str]) -> np.ndarray:
    if not chunks:
        raise ValueError("embedding할 chunk가 없습니다.")

    model = get_model()

    output = model.encode(
        chunks,
        batch_size=4,
        max_length=512,
        return_dense=True,
        return_sparse=False,
        return_colbert_vecs=False,
    )

    embeddings = np.asarray(
        output["dense_vecs"],
        dtype=np.float32,
    )

    if embeddings.shape != (len(chunks), EXPECTED_DIM):
        raise RuntimeError(
            f"잘못된 embedding shape: {embeddings.shape}"
        )

    if not np.isfinite(embeddings).all():
        raise RuntimeError("embedding에 NaN 또는 Inf가 있습니다.")

    norms = np.linalg.norm(embeddings, axis=1)

    if np.any(norms <= 0):
        raise RuntimeError("0 vector가 생성되었습니다.")

    return embeddings


def verify_document_owner(
    document_id: int,
    owner_user_id: int,
) -> None:
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT owner_user_id
                FROM documents
                WHERE id = %s
                """,
                (document_id,),
            )
            row = cur.fetchone()

    if row is None:
        raise ValueError(
            f"존재하지 않는 document_id입니다: {document_id}"
        )

    if row[0] != owner_user_id:
        raise PermissionError(
            "문서 소유자가 일치하지 않습니다."
        )


def save_chunks(
    document_id: int,
    owner_user_id: int,
    chunks: list[str],
    embeddings: np.ndarray,
) -> None:
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT owner_user_id
                FROM documents
                WHERE id = %s
                FOR UPDATE
                """,
                (document_id,),
            )
            row = cur.fetchone()

            if row is None:
                raise ValueError(
                    f"존재하지 않는 document_id입니다: {document_id}"
                )

            if row[0] != owner_user_id:
                raise PermissionError(
                    "문서 소유자가 일치하지 않습니다."
                )

            # 재인덱싱 시 기존 chunk 제거
            cur.execute(
                """
                DELETE FROM document_chunks
                WHERE document_id = %s
                """,
                (document_id,),
            )

            sql = """
                INSERT INTO document_chunks (
                    document_id,
                    chunk_index,
                    content,
                    embedding
                )
                VALUES (%s, %s, %s, %s::vector)
            """

            rows = []

            for chunk_index, content in enumerate(chunks):
                rows.append(
                    (
                        document_id,
                        chunk_index,
                        content,
                        vector_literal(
                            embeddings[chunk_index]
                        ),
                    )
                )

            cur.executemany(sql, rows)

        conn.commit()


def index_document(
    document_id: int,
    owner_user_id: int,
    file_bytes: bytes,
    filename: str | None = None,
    file_type: str | None = None,
) -> dict:
    resolved_file_type = resolve_file_type(
        filename=filename,
        file_type=file_type,
    )

    verify_document_owner(
        document_id=document_id,
        owner_user_id=owner_user_id,
    )

    text = extract_text(
        file_bytes=file_bytes,
        file_type=resolved_file_type,
    )

    chunks = chunk_text(
        text,
        min_chars=300,
        target_chars=400,
        max_chars=500,
    )

    if not chunks:
        raise ValueError("chunking 결과가 비어 있습니다.")

    embeddings = embed_chunks(chunks)

    save_chunks(
        document_id=document_id,
        owner_user_id=owner_user_id,
        chunks=chunks,
        embeddings=embeddings,
    )

    return {
        "document_id": document_id,
        "owner_user_id": owner_user_id,
        "file_type": resolved_file_type,
        "text_chars": len(text),
        "chunk_count": len(chunks),
        "embedding_dimension": EXPECTED_DIM,
        "status": "indexed",
    }
